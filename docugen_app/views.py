from django.shortcuts import render
from docx import Document
from docx.shared import Pt
from datetime import date
from django.shortcuts import redirect
from .models import Proektant
# Create your views here.


def index(request):
    template = 'home.html'
    context = {}
    return render(request, template, context)


def kompleksen_doklad(request):
    if request.method == 'GET':
        # form = KompleksenDokladForm()
        proektants = Proektant.objects.all()
        template = 'pages/kompleksen.html'

        context = {
            # 'form': form
            'proektants': proektants
        }

        return render(request, template, context)
    
    if request.method == 'POST':
        form_data = request.POST
        document = Document()
        title = document.add_paragraph().add_run('АРЕА ПРОЕКТ ЕООД, гр. Разград, ул. „Кирил и Методий“ №3, ет2, офис10')
        title.font.size = Pt(14)
        title.bold = True
        document.add_paragraph('тел/факс: 084 620000; 0893 316563, e-mal: area_proekt@abv.bg')
        doc_num = document.add_paragraph().add_run('Изх. №  {docnum} / {date} '.format(docnum=form_data['docnum'], date=date.today()))
        doc_num.font.size = Pt(11)
        vazlozhitel_title = document.add_paragraph().add_run('До ВЪЗЛОЖИТЕЛЯ :') 
        vazlozhitel_title.font.size = Pt(14)
        vazlozhitel_title.bold = True
        vazlozhitel = document.add_paragraph().add_run(form_data['build'])
        vazlozhitel.font.size = Pt(14)
        vazlozhitel.bold = True
        
        document.save('demo.docx')

        return redirect('/')


def okonchatelen_doklad(request):
    if request.method == 'GET':
        # form = OkonchatelenDokladForm()

        template = 'pages/okonchatelen.html'

        context = {
            # 'form': form
        }

        return render(request, template, context)